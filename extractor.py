from tkinter import Tk, filedialog
from typing import List, Union, Dict
# from utils.loader import load_json
# import pycountry
import re
import docx

class ValueExtractor:
    def __init__(self):
        self.week_patterns = [
            # Week of the year, e.g., Week 35
            r'Week\s*\d{1,2}',
            # Weekday, e.g., Monday, Tuesday, etc.
            r'(?:Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday)',

            r'第\d{1,2}\s*?周',  # matches 第1 周, 第12周, etc.
            r'(?:星期一|星期二|星期三|星期四|星期五|星期六|星期日)',  # matches 星期一, 星期二, etc.
        ]


        self.date_patterns = [
            # English dates
            r'\d{1,2}\s*?(?:January|February|March|April|May|June|July|August|September|October|November|December)\s*?\d{4}',
            r'\d{1,2}\s*?(?:January|February|March|April|May|June|July|August|September|October|November|December)',
            r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s*?\d{4}',
            # Chinese dates
            # 2022年7月22日 or 2022 年 7 月 22 日
            r'\d{4}\s*?年\s*?\d{1,2}\s*?月\s*?\d{1,2}\s*?日',
            # 2022年7月 or 2022 年 7 月
            r'\d{4}\s*?年\s*?\d{1,2}\s*?月',
            r'\d{1,2}\s*?月\s*?\d{1,2}\s*?日',                  # 3月7日 or 3 月 7 日
            r'\d{1,2}\s*?月'                                  # 3月 or 3 月
        ]
        
        self.percentage_patterns = [
            # Match number followed by % (e.g., 50%, 5.5%, 0.123%)
            r'\d+(?:\.\d+)?\s*%',
            # Chinese percentage pattern
            r'\d+(?:\.\d+)?\s*百分之'
        ]

        self.numerical_patterns = [
            # English Numerical Units Pattern
            r'\d+\s*(?:hundred|thousand|million|billion)s?',
            # Chinese Numerical Units Pattern
            # r'(?:一|十|百|千)?\d+\s*(?:億|萬|百萬|千萬)?',
            # r'\b(?:一|十|百|千)?\d+\s*(?:億|萬|百萬|千萬)?\b'

        ]
        self.time_patterns = [
            # 24-hour time format (e.g., 9:30, 15:45, 00:00)
            r'\d{1,2}(?::\d{2})',

            # matches 9.00 am, 9.00am, 9:00PM, 6p.m., 6:00p.m., etc.
            r'\d{1,2}(?:[.:]\d{2})?\s?[apAP]\.?m\.?',
            
            # Noon and midnight (e.g., 12 noon, 12 midnight)
            r'12\s?(?:noon|midnight)',
            
            # 24-hour time format in Chinese (e.g., 9 時, 15 時 30 分, 00 時)
            r'\d{1,2}\s?時(?:\s?\d{1,2}\s?分)?',
            
            # 12-hour time format in Chinese (e.g., 9 時 30 分 AM, 3 時 15 分 PM)
            r'\d{1,2}\s?時\s?\d{1,2}\s?分\s?[APap]\.?m\.?',
            
        ]
        
        self.quantity_patterns = [
            # Percentage
            r'\d+(?:\.\d{1,2})?\s*%',               # matches 1 %, 1%, etc.

            # US Dollar
            r'US\$\d+(?:\.\d{1,2})?',   # matches US$1.00, US$1, US$0.50, etc.
            r'\$\d+(?:\.\d{1,2})?',    # matches $1.00, $1, etc.
            r'\s+\d+\s?USD',               # matches 1 USD, 1USD, etc.

            # Euro
            r'€\d+(?:\.\d{1,2})?',     # matches €1.00, €1, etc.
            r'\s+\d+\s?EUR',               # matches 1 EUR, 1EUR, etc.

            # Chinese Yuan/Renminbi
            r'¥\d+(?:\.\d{1,2})?',     # matches ¥1.00, ¥1, etc.
            r'\s+\d+\s?CNY',               # matches 1 CNY, 1CNY, etc.
            r'\s+\d+\s?RMB',               # matches 1 RMB, 1RMB, etc.

            # Singapore Dollar
            r'S\$\d+(?:\.\d{1,2})?',   # matches S$1.00, S$1, etc.
            r'\s+\d+\s?SGD',               # matches 1 SGD, 1SGD, etc.

            # Chinese expressions
            r'\d+\s?美元',              # matches 1 美元 for US dollar
            r'\d+\s?元',                # matches 1 元 for RMB
            r'\d+\s?新加坡元',          # matches 1 新加坡元 for Singapore dollar
            r'\d+\s?欧元',              # matches 1 欧元 for Euro
        ]

        self.word_to_number = {
            'one': '1',
            'two': '2',
            'three': '3',
            'four': '4',
            'five': '5',
            'six': '6',
            'seven': '7',
            'eight': '8',
            'nine': '9',
            'ten': '10',
            'eleven': '11',
            'twelve': '12',
            'thirteen': '13',
            'fourteen': '14',
            'fifteen': '15',
            'sixteen': '16',
            'seventeen': '17',
            'eighteen': '18',
            'nineteen': '19',
            'twenty': '20',
            'third': '3',
            'thirds': '3',
            'quarter': '4',
            'quarters': '4',
            'half': '2',
        }

        def get_country_list():
            # Get a list of all countries in English, their corresponding names in Chinese (Simplified), and their short forms.
            countries = {}
            chinese_country_map = load_json("resources/country_map.json")
            countries.update(chinese_country_map)
            for country in pycountry.countries:
                countries[country.name.lower()] = country.name
                countries[country.alpha_2] = country.name

                try:
                    cn_name = country.name_translations['zh-Hans']
                    countries[cn_name] = country.name
                except (KeyError, AttributeError):
                    pass
            return countries

        # self.countries = get_country_list()

    def _chinese_to_number(self, chinese_num: str) -> int:
        chinese_numerals = {
            '零': 0, '一': 1, '二': 2, '三': 3, '四': 4,
            '五': 5, '六': 6, '七': 7, '八': 8, '九': 9
        }
        unit_positions = {
            '十': 10, '百': 100, '千': 1000, '万': 10000
        }
        val = 0
        unit_val = 1
        for char in reversed(chinese_num):
            if char in unit_positions:
                unit_val = unit_positions[char]
            else:
                val += chinese_numerals[char] * unit_val
                if unit_val > 1:
                    unit_val = 1
        return val

    def extract_date_values_from_text(self, s: str) -> List[str]:
        """
        Extracts date values from the given string.

        Args:
        - s (str): The input string.

        Returns:
        - List[str]: List of extracted date values.
        """

        date_values = []

        # Iterating in the order of the patterns
        for pattern in self.date_patterns:
            matches = re.findall(pattern, s)
            if matches:
                # Extend date_values with the new matches
                date_values.extend(matches)

                # Remove these matches from the string to avoid extracting sub-patterns again
                for match in matches:
                    # replace only the first occurrence
                    s = s.replace(match, '', 1)

        return date_values

    def extract_numerical_values_from_text(self, s: str) -> List[str]:
        """
        Extracts numerical values from the given string, excluding those that are part of a date.

        Args:
        - s (str): The input string.

        Returns:
        - List[str]: List of extracted numerical values.
        """

        # Extract date values
        date_values = self.extract_date_values_from_text(s)

        # Remove extracted date values from the string
        for date in date_values:
            s = s.replace(date, '')

        # Extract fraction values
        fraction_values = self.extract_fractions_from_text(s)

        # Remove extracted fraction values from the string
        for fraction in fraction_values:
            s = s.replace(fraction, '')

        # Extract time values
        time_values = self.extract_time_values_from_text(s)

        # Remove extracted time values from the string
        for time in time_values:
            s = s.replace(time, '')

        # Extract quantity values
        quantity_values = self.extract_quantity_values_from_text(s)

        # Remove extracted quantity values from the string
        for quantity in quantity_values:
            s = s.replace(quantity, '')

        # Find chinese order words and replace into Arabic number
        replace_keys_ch = {
            '零級': '0',
            '一級': '1',
            '二級': '2',
            '三級': '3',
            '四級': '4',
            '五級': '5',
            '六級': '6',
            '七級': '7',
            '八級': '8',
            '九級': '9',
            '第零': '0',
            '第一': '1',
            '第二': '2',
            '第三': '3',
            '第四': '4',
            '第五': '5',
            '第六': '6',
            '第七': '7',
            '第八': '8',
            '第九': '9',
        }
        # Remove extracted order values from the string
        for key in replace_keys_ch:
            s = s.replace(key, replace_keys_ch[key])

        # Extract numbers, excluding those within date patterns
        matches = re.findall(r'(\d+(?:,\d{3})*(?:\.\d+)?)', s)
        result = [match.replace(",", "") for match in matches]
        return result

    def extract_time_values_from_text(self, s: str) -> List[str]:
        """
        Extracts time values from the given string, excluding those that are part of a date.

        Args:
        - s (str): The input string.

        Returns:
        - List[str]: List of extracted times.
        """

        time_values = []

        # Iterating in the order of the patterns
        for pattern in self.time_patterns:
            matches = re.findall(pattern, s)
            if matches:
                # Extend time_values with the new matches
                time_values.extend(matches)

                # Remove these matches from the string to avoid extracting sub-patterns again
                for match in matches:
                    # replace only the first occurrence
                    s = s.replace(match, '', 1)

        return time_values

    def extract_quantity_values_from_text(self, s: str) -> List[str]:
        """
        Extracts quantity values from the given string, excluding those that are part of a date.

        Args:
        - s (str): The input string.

        Returns:
        - List[str]: List of extracted quantitys.
        """

        quantity_values = []

        # Iterating in the order of the patterns
        for pattern in self.quantity_patterns:
            matches = re.findall(pattern, s)
            if matches:
                # Extend quantity_values with the new matches
                quantity_values.extend(matches)

                # Remove these matches from the string to avoid extracting sub-patterns again
                for match in matches:
                    # replace only the first occurrence
                    s = s.replace(match, '', 1)

        return quantity_values

    def extract_fractions_from_text(self, text: str) -> List[str]:
        """
        Extracts fractions from both English and Chinese text.

        Args:
        - text (str): The input string.

        Returns:
        - List[str]: List of extracted fractions.
        """
        extracted_fractions = []

        try:
            # For English fractions
            extracted_fractions.extend(re.findall(r'\b\d+/\d+\b', text))

            # For worded English fractions
            worded_english_fractions_raw = re.findall(
                r'\b(one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|fifteen|sixteen|seventeen|eighteen|nineteen|twenty)-(\w+)\b', text)
            extracted_fractions.extend(
                f"{self.word_to_number.get(numerator, 'unknown')}/{self.word_to_number.get(denominator, 'unknown')}"
                for numerator, denominator in worded_english_fractions_raw if numerator in self.word_to_number and denominator in self.word_to_number
            )

            # For Chinese fractions
            chinese_fractions_raw = re.findall(
                r'([零一二三四五六七八九十百千万]+)分之([零一二三四五六七八九十百千万]+)', text)
            extracted_fractions.extend(
                f"{self._chinese_to_number(numerator)}/{self._chinese_to_number(denominator)}"
                for denominator, numerator in chinese_fractions_raw
            )
        except Exception as e:
            # You can log the error if needed or take any other corrective actions.
            print(f"Error encountered: {e}")

        return extracted_fractions

    def extract_country_names_from_text(self, text: str, lang: str = "en") -> List[str]:
        found_countries = set()

        # Define the patterns to remove as a regular expression
        patterns_to_remove = [r"\.", r"LAP", r"ESG"]
        combined_pattern = "|".join(patterns_to_remove)

        cleaned_text = re.sub(combined_pattern, "", text)

        for country in self.countries:
            if len(country) == 2:  # If it's an alpha-2 code, it should be uppercase
                if lang == "en":
                    pattern = rf"\b{re.escape(country)}\b"
                else:
                    pattern = rf"{re.escape(country)}"
                if re.search(pattern, cleaned_text):
                    found_countries.add(self.countries[country])
            else:
                if lang == "en":
                    pattern = rf"\b{re.escape(country)}\b"
                else:
                    pattern = rf"{re.escape(country)}"
                if re.search(pattern, cleaned_text, re.IGNORECASE):
                    found_countries.add(self.countries[country])

        return list(found_countries)

# Create an instance of the ValueExtractor class
extractor = ValueExtractor()
# Function to extract and print values from text
def extract_and_print_values(text):
    # Extract date patterns
    date_matches = []
    for pattern in extractor.date_patterns:
        date_matches.extend(re.findall(pattern, text))

    # Extract time patterns
    time_matches = []
    for pattern in extractor.time_patterns:
        time_matches.extend(re.findall(pattern, text))

    # Extract quantity patterns
    quantity_matches = []
    for pattern in extractor.quantity_patterns:
        quantity_matches.extend(re.findall(pattern, text))

        
    week_matches = []
    for pattern in extractor.week_patterns:
        week_matches.extend(re.findall(pattern, text))
        
    percentage_matches = []
    for pattern in extractor.percentage_patterns:
        percentage_matches.extend(re.findall(pattern, text))
        
    numerical_matches = []
    for pattern in extractor.numerical_patterns:
        numerical_matches.extend(re.findall(pattern, text))

    # Extracted values
    extracted_values = {
        # "Dates": date_matches,
        # "Times": time_matches,
        # "Quantities": quantity_matches,
        # "Weeks": week_matches,
        # "Percentages": percentage_matches,
        "Numericals": numerical_matches,
    }

    # Print the extracted values
    for category, values in extracted_values.items():
        if values:
            print(f"{category}:")
            for value in values:
                print(f"- {value}")

# Function to extract and process text from paragraphs
def extract_text_from_paragraphs(doc):
    doc_text = ""
    for paragraph in doc.paragraphs:
        doc_text += paragraph.text + "\n"
    extract_and_print_values(doc_text)

# Function to extract and process text from tables
def extract_text_from_tables(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                extract_and_print_values(cell_text)

# Function to extract and process hyperlinks
def extract_hyperlinks(doc):
    for hyperlink in doc.element.xpath('//w:hyperlink'):
        link_text = hyperlink.xpath('./w:r/w:t')
        if link_text:
            link_text = link_text[0].text
            print(f"Hyperlink: {link_text}")

# Open a file dialog to select a docx file
root = Tk()
root.withdraw()  # Hide the main window
file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])

if file_path:
    # Load the docx file
    doc = docx.Document(file_path)

    # Extract and process text from paragraphs, tables, and hyperlinks
    extract_text_from_paragraphs(doc)
    extract_text_from_tables(doc)
    # extract_hyperlinks(doc)
else:
    print("No file selected.")